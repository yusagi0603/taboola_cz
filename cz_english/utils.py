import json
from pathlib import Path

from typing_extensions import override

from config import (
    OPENAI_API_KEY,
    ASSISTANT_NAME,
    ASSISTANT_DESCRIPTION,
    ASSISTANT_MODEL,
    CORRECT_PASSWORD,
    ASSISTANT_ID,
)

from openai import AssistantEventHandler

import json

from docx import Document
    
from logger import app_logger
from client import llm_client

import gradio as gr

ARTICLE_GENERATION_PATH = Path(__file__).parent / "prompt" / "article_generation.jinja"
ARTICLE_REWRITE_PATH = Path(__file__).parent / "prompt" / "article_rewrite.jinja"
QUESTION_FORMAT_PATH = Path(__file__).parent / "prompt" / "question_format.jinja"

with open(ARTICLE_GENERATION_PATH, 'r', encoding='utf-8') as f:
    ARTICLE_GENERATION = f.read()

with open(ARTICLE_REWRITE_PATH, 'r', encoding='utf-8') as f:
    ARTICLE_REWRITE = f.read()

with open(QUESTION_FORMAT_PATH, 'r', encoding='utf-8') as f:
    QUESTION_FORMAT = f.read()


# spent 4m 50s downloading all 175 files
# def embed_from_drive(folder_id):
#   # auth.authenticate_user()
#   gauth = GoogleAuth()
#   gauth.credentials = GoogleCredentials.get_application_default()
#   drive = GoogleDrive(gauth)

#   # Get all files in '定稿專案' folder: https://drive.google.com/drive/folders/1dlsf5BNjNczzUYKPZvYXd2mLW21QCLUK?usp=drive_link
#   file_list = drive.ListFile({'q': f"'{folder_id}' in parents and trashed=false"}).GetList()

#   # Download files to local (`/content/`), since file_streams don't recieve google docs
#   local_file_paths = []
#   for file1 in file_list:
#       print('Processing file title: %s, id: %s' % (file1['title'], file1['id']))
#       local_path = f"/content/{file1['title']}.docx"

#       if 'exportLinks' in file1:
#           if 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' in file1['exportLinks']:
#               # update type if needed (application/vnd.openxmlformats-officedocument.wordprocessingml.document == .docx)
#               export_url = file1['exportLinks']['application/vnd.openxmlformats-officedocument.wordprocessingml.document']
#               print(f"Downloading as Word document: {file1['title']}")
#               downloaded_file = drive.CreateFile({'id': file1['id']})
#               downloaded_file.GetContentFile(local_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#               local_file_paths.append(local_path)
#           else:
#               print(f"No Word export available for: {file1['title']}")
#       else:
#           print(f"Skipping non-Google Docs file: {file1['title']}")

#   for path in local_file_paths:
#       print(f"Downloaded file: {path}")

#   file_streams = [open(path, "rb") for path in local_file_paths]
#   return file_streams


# Embed files (downloaded from drive folder)
# def get_vector_store_id(file_streams):
#   vector_store = client.beta.vector_stores.create(name=VECTOR_STORE_NAME)

# #   # spent 51s batching all 175 files
# #   file_batch = client.beta.vector_stores.file_batches.upload_and_poll(
# #     vector_store_id=vector_store.id, files=file_streams
# #   )

# #   print("file_batch status",file_batch.status)
# #   print("file_counts",file_batch.file_counts)

#   return vector_store.id



def check_password(input_password):
    if input_password == CORRECT_PASSWORD:
        return gr.update(visible=False), gr.update(visible=True), ""
    else:
        return gr.update(visible=True), gr.update(visible=False), gr.update(value="Wrong Password. Please Retry. hint: channel name", visible=True)


def _call_llm_with_prompt(system_prompt, user_prompt, response_format=None):
    msg_lst = []
    for role, prompt in zip(["system", "user"], [system_prompt, user_prompt]):
        if prompt is not None:
            msg_lst.append({"role": role, "content": prompt})

    response = llm_client.chat.completions.create(
        model=ASSISTANT_MODEL,
        messages=msg_lst,
        response_format=response_format
    )
    generated_article = response.choices[0].message.content
    return generated_article    


def call_llm_to_generate_article(
        grade_values, unit_values, topic_values, grammar_values, input_article_value
    ):

    if input_article_value:
        user_prompt = ARTICLE_REWRITE.format(
            grade_values=grade_values,
            topic_values=topic_values,
            grammar_values=grammar_values,
            vocabulary_values=unit_values,
            input_article_value=input_article_value
        )
        app_logger.info("Rewriting article with input")
    else:
        user_prompt = ARTICLE_GENERATION.format(
            grade_values=grade_values,
            topic_values=topic_values,
            grammar_values=grammar_values,
            vocabulary_values=unit_values
        )
        app_logger.info("Generating article without input")

    generated_article = _call_llm_with_prompt(
        system_prompt=None, 
        user_prompt=user_prompt,
        response_format=None
    )
    return generated_article


def call_llm_to_generate_question(question_type):

    system_prompt, user_prompt = "You are a question generator for English exam in Taiwan.", \
        f"Please generate a question based on the question type: {question_type}"

    generated_question = _call_llm_with_prompt(
        system_prompt=system_prompt,
        user_prompt=user_prompt,
        response_format=QUESTION_FORMAT
    )
    return generated_question


def create_google_doc(self, title, content):
    """Create a Google Doc with the given title and content."""
    try:
        # Import necessary libraries for Google Docs API
        from googleapiclient.discovery import build
        from google.oauth2 import service_account
        
        # Set up credentials and service
        SCOPES = ['https://www.googleapis.com/auth/documents']
        SERVICE_ACCOUNT_FILE = 'credentials.json'  # Path to your service account credentials
        
        credentials = service_account.Credentials.from_service_account_file(
            SERVICE_ACCOUNT_FILE, scopes=SCOPES)
        docs_service = build('docs', 'v1', credentials=credentials)
        
        # Create a new document
        document = {
            'title': title
        }
        doc = docs_service.documents().create(body=document).execute()
        document_id = doc.get('documentId')
        
        # Insert content into the document
        requests = [
            {
                'insertText': {
                    'location': {
                        'index': 1
                    },
                    'text': content
                }
            }
        ]
        
        docs_service.documents().batchUpdate(
            documentId=document_id,
            body={'requests': requests}
        ).execute()
        
        # Get the document URL
        doc_url = f"https://docs.google.com/document/d/{document_id}/edit"
        return doc_url
        
    except Exception as e:
        app_logger.error(f"Error creating Google Doc: {str(e)}")
        return None


def generate_docx_file(doc_file_name, insert_doc_info):
    
    # # Create a Word document
    doc = Document()
    doc.add_heading(doc_file_name, 0)
    
    # # Add each question to the document
    for content_type, content in insert_doc_info:
        doc.add_heading(f"{content_type}", level=1)
        doc.add_paragraph(content)
        doc.add_paragraph("")  # Add some spacing

    doc.save(doc_file_name)

    return doc_file_name


from datetime import datetime
from docx import Document
from exam_maker.logger import app_logger

class ExamPaperHandler:
    def __init__(self):
        pass

    def generate_final_exam_doc(self, article_content, problem_list):
        """
        Generate a final exam paper from the article content and problem list.
        
        Args:
            article_content (str): The main article content
            problem_list (list): List of tuples containing (problem_type, problem_content)
            
        Returns:
            tuple: (doc_file_name, visibility_update)
        """
        # Compose the full exam content
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc_file_name = f"英文考題 - {timestamp}"
        
        # Return the exam content as a downloadable document
        question_info_tuple = [("文章", article_content)]
        for problem_type, problem_content in problem_list:
            question_info_tuple.append((problem_type, problem_content))
            
        doc_file_name = self.generate_docx_file(
            doc_file_name,
            question_info_tuple
        )
        
        return doc_file_name, True  # True for visibility update

    def generate_docx_file(self, doc_file_name, insert_doc_info):
        """
        Generate a Word document with the given content.
        
        Args:
            doc_file_name (str): Name of the document file
            insert_doc_info (list): List of tuples containing (content_type, content)
            
        Returns:
            str: The name of the generated file
        """
        # Create a Word document
        doc = Document()
        doc.add_heading(doc_file_name, 0)
        
        # Add each question to the document
        for content_type, content in insert_doc_info:
            doc.add_heading(f"{content_type}", level=1)
            doc.add_paragraph(content)
            doc.add_paragraph("")  # Add some spacing

        doc.save(doc_file_name)
        return doc_file_name

    def create_google_doc(self, title, content):
        """
        Create a Google Doc with the given title and content.
        
        Args:
            title (str): The title of the document
            content (str): The content to be inserted
            
        Returns:
            str: The URL of the created Google Doc, or None if creation failed
        """
        try:
            # Import necessary libraries for Google Docs API
            from googleapiclient.discovery import build
            from google.oauth2 import service_account
            
            # Set up credentials and service
            SCOPES = ['https://www.googleapis.com/auth/documents']
            SERVICE_ACCOUNT_FILE = 'credentials.json'  # Path to your service account credentials
            
            credentials = service_account.Credentials.from_service_account_file(
                SERVICE_ACCOUNT_FILE, scopes=SCOPES)
            docs_service = build('docs', 'v1', credentials=credentials)
            
            # Create a new document
            document = {
                'title': title
            }
            doc = docs_service.documents().create(body=document).execute()
            document_id = doc.get('documentId')
            
            # Insert content into the document
            requests = [
                {
                    'insertText': {
                        'location': {
                            'index': 1
                        },
                        'text': content
                    }
                }
            ]
            
            docs_service.documents().batchUpdate(
                documentId=document_id,
                body={'requests': requests}
            ).execute()
            
            # Get the document URL
            doc_url = f"https://docs.google.com/document/d/{document_id}/edit"
            return doc_url
            
        except Exception as e:
            app_logger.error(f"Error creating Google Doc: {str(e)}")
            return None 